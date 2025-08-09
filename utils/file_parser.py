import streamlit as st
import tempfile
import os
from typing import Union, Optional
from pathlib import Path

class FileParser:
    """Utility class for parsing different file formats"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._parse_pdf,
            'txt': self._parse_txt,
            'doc': self._parse_doc,
            'docx': self._parse_docx,
            'py': self._parse_python
        }
    
    def parse_file(self, uploaded_file) -> str:
        """Parse uploaded file and return its content as string"""
        if uploaded_file is None:
            raise ValueError("No file provided")
        
        # Get file extension
        file_extension = self._get_file_extension(uploaded_file.name)
        
        # Check if format is supported
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Parse file based on its type
        try:
            return self.supported_formats[file_extension](uploaded_file)
        except Exception as e:
            raise Exception(f"Error parsing {uploaded_file.name}: {str(e)}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename"""
        return Path(filename).suffix.lower().lstrip('.')
    
    def _parse_txt(self, uploaded_file) -> str:
        """Parse TXT file"""
        try:
            content = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Try with different encoding
            uploaded_file.seek(0)
            return content.decode('latin-1')
    
    def _parse_python(self, uploaded_file) -> str:
        """Parse Python file"""
        return self._parse_txt(uploaded_file)
    
    def _parse_pdf(self, uploaded_file) -> str:
        """Parse PDF file using PyMuPDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Install with: pip install PyMuPDF")
        
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name
            
            # Parse PDF
            doc = fitz.open(tmp_file_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            # Reset file pointer
            uploaded_file.seek(0)
            
            return text_content.strip()
            
        except Exception as e:
            # Clean up on error
            if 'tmp_file_path' in locals():
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_doc(self, uploaded_file) -> str:
        """Parse DOC file using python-docx"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name
            
            # Parse DOC
            doc = Document(tmp_file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            # Reset file pointer
            uploaded_file.seek(0)
            
            return text_content.strip()
            
        except Exception as e:
            # Clean up on error
            if 'tmp_file_path' in locals():
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
            raise Exception(f"Error parsing DOC: {str(e)}")
    
    def _parse_docx(self, uploaded_file) -> str:
        """Parse DOCX file using python-docx"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name
            
            # Parse DOCX
            doc = Document(tmp_file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            # Reset file pointer
            uploaded_file.seek(0)
            
            return text_content.strip()
            
        except Exception as e:
            # Clean up on error
            if 'tmp_file_path' in locals():
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def validate_file_size(self, uploaded_file, max_size_mb: int = 10) -> bool:
        """Validate file size"""
        if uploaded_file is None:
            return False
        
        file_size_mb = uploaded_file.size / (1024 * 1024)
        return file_size_mb <= max_size_mb
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get file information"""
        if uploaded_file is None:
            return {}
        
        return {
            'name': uploaded_file.name,
            'size_mb': round(uploaded_file.size / (1024 * 1024), 2),
            'type': uploaded_file.type,
            'extension': self._get_file_extension(uploaded_file.name)
        } 